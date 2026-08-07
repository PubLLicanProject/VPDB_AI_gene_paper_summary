"""
supplementary_helpers.py

Dynamically fetch and parse the supplementary files of an open-access PMC article, then
extract the parts relevant to a given gene, so they can augment the main-text summary input.

Retrieval (nothing is written to disk):
    PMID  --(NCBI ID Converter)-->  PMCID
    PMCID --(Europe PMC)-->  /supplementaryFiles  -> ZIP archive (held in memory)
    PMCID --(Europe PMC)-->  /fullTextXML         -> JATS, for each file's label + caption

Parsing by format:
    .pdf                 -> pdfplumber (page text + extracted tables)
    .xlsx / .xls         -> pandas (openpyxl / xlrd), one section per sheet
    .csv/.tsv/.tab/.txt  -> encoding sniff (charset_normalizer) + delimiter sniff (csv.Sniffer)
    .docx                -> python-docx (paragraphs + tables)
    .xml / .html         -> BeautifulSoup text
    images               -> deferred (filename + caption placeholder only)

Gene-aware filtering keeps only the table rows / text lines that mention the gene id or an
alias, plus each file's header/caption, so the block handed to the LLM stays small and
relevant. Every failure degrades to "" — this must never block the pipeline.

Public entry point:
    get_supplementary_text(pmid, gene_id, aliases=None, host_db=None, *, caps=None) -> str
"""

import io
import os
import re
import csv
import time
import random
import zipfile
import threading
import requests
from collections import OrderedDict
from typing import List, Dict, Optional

# --- Endpoints --------------------------------------------------------------
ID_CONVERTER_URL = "https://www.ncbi.nlm.nih.gov/pmc/utils/idconv/v1.0/"
EPMC_BASE = "https://www.ebi.ac.uk/europepmc/webservices/rest"

# --- Fetch robustness (mirrors pubmed_helpers) ------------------------------
HTTP_TIMEOUT = 60
FETCH_MAX_RETRIES = 4
FETCH_BACKOFF_BASE = 1.5
FETCH_BACKOFF_CAP = 30.0
RETRYABLE_STATUS = {429, 500, 502, 503, 504}

# NCBI politeness — the ID Converter is a genuine NCBI utility that honours these.
NCBI_TOOL = os.getenv("NCBI_TOOL", "VPDB_PD_generator")
NCBI_EMAIL = os.getenv("NCBI_EMAIL", "")
NCBI_API_KEY = os.getenv("NCBI_API_KEY", "")

XLINK = "{http://www.w3.org/1999/xlink}href"
IMAGE_EXTS = {".gif", ".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".svg", ".eps"}

# --- Default size caps (override via the `caps` argument / config) -----------
DEFAULT_CAPS = {
    "max_zip_bytes": 60 * 1024 * 1024,        # compressed-download ceiling (streamed, enforced live)
    "max_file_bytes": 25 * 1024 * 1024,       # per-file UNCOMPRESSED read ceiling (via _read_capped)
    "max_total_extract_bytes": 150 * 1024 * 1024,  # cumulative uncompressed bytes read (incl. nested)
    "max_files": 200,                         # max entries processed (incl. nested)
    "max_zip_depth": 1,                        # nested-zip recursion depth
    "max_matched_rows": 40,                   # per-file rows/lines kept after gene filter
    "max_always": 6,                          # header/label lines kept per file
    "preview_rows": 2,                        # rows shown when a file has no gene hit
    "max_file_chars": 8000,                   # per-file output ceiling
    "total_char_budget": 24000,               # overall supplementary block ceiling
}

# Small bounded caches so a paper's supplement is downloaded/parsed once, not per gene
# (a paper with N genes otherwise re-fetches the same zip N times).
_ZIP_CACHE = OrderedDict()      # pmcid -> zip bytes | None
_CAPTION_CACHE = OrderedDict()  # pmcid -> {filename: caption}
_PARSED_CACHE = OrderedDict()   # pmcid -> {"files":[(name,parsed)], "image_count", "other_files"} | None
_CACHE_MAX = 24                 # >= worker count so concurrent papers don't evict each other
_cache_lock = threading.Lock()  # guards the OrderedDicts under paper-level threading


def _cache_put(cache, key, value):
    with _cache_lock:
        cache[key] = value
        cache.move_to_end(key)
        while len(cache) > _CACHE_MAX:
            cache.popitem(last=False)
    return value


# ---------------------------------------------------------------------------
# HTTP
# ---------------------------------------------------------------------------
def _http_get(url, *, params=None, timeout=HTTP_TIMEOUT, max_retries=FETCH_MAX_RETRIES):
    """GET with retry/backoff on transient errors. Returns a 200 Response or raises."""
    last = "unknown error"
    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(url, params=params, timeout=timeout)
        except (requests.Timeout, requests.ConnectionError) as e:
            last = f"{type(e).__name__}: {e}"
        else:
            if resp.status_code == 200:
                return resp
            if resp.status_code in RETRYABLE_STATUS:
                last = f"HTTP {resp.status_code}"
                retry_after = resp.headers.get("Retry-After")
                if retry_after:
                    try:
                        time.sleep(min(float(retry_after), FETCH_BACKOFF_CAP))
                        continue
                    except ValueError:
                        pass
            else:
                raise RuntimeError(f"GET {url} -> HTTP {resp.status_code}")
        if attempt < max_retries:
            time.sleep(min(FETCH_BACKOFF_CAP, FETCH_BACKOFF_BASE * (2 ** attempt)) * (0.5 + random.random()))
    raise RuntimeError(f"GET {url} failed after {max_retries + 1} attempts: {last}")


def _ncbi_params(extra=None):
    params = dict(extra or {})
    if NCBI_TOOL:
        params["tool"] = NCBI_TOOL
    if NCBI_EMAIL:
        params["email"] = NCBI_EMAIL
    if NCBI_API_KEY:
        params["api_key"] = NCBI_API_KEY
    return params


def _download_capped(url, max_bytes, *, timeout=HTTP_TIMEOUT, max_retries=FETCH_MAX_RETRIES):
    """
    Stream a GET with retry/backoff and a hard size ceiling. Returns bytes, or None if the
    body exceeds `max_bytes` (declared or actual) or the request fails. Never buffers an
    over-cap body: aborts on a too-large Content-Length and mid-stream once the cap is hit.
    """
    last = "unknown error"
    for attempt in range(max_retries + 1):
        resp = None
        try:
            resp = requests.get(url, timeout=timeout, stream=True)
            status = resp.status_code
            if status == 200:
                clen = resp.headers.get("Content-Length")
                if clen and clen.isdigit() and int(clen) > max_bytes:
                    return None  # declared too big — do not download
                buf = bytearray()
                for chunk in resp.iter_content(chunk_size=65536):
                    if not chunk:
                        continue
                    buf.extend(chunk)
                    if len(buf) > max_bytes:
                        return None  # exceeded cap mid-stream — abort
                return bytes(buf)
            if status in RETRYABLE_STATUS:
                last = f"HTTP {status}"
                retry_after = resp.headers.get("Retry-After")
                if retry_after:
                    try:
                        time.sleep(min(float(retry_after), FETCH_BACKOFF_CAP))
                        continue
                    except ValueError:
                        pass
            else:
                return None  # non-retryable (e.g. 404 = no supplementary files)
        except (requests.Timeout, requests.ConnectionError) as e:
            last = f"{type(e).__name__}: {e}"
        finally:
            if resp is not None:
                resp.close()
        if attempt < max_retries:
            time.sleep(min(FETCH_BACKOFF_CAP, FETCH_BACKOFF_BASE * (2 ** attempt)) * (0.5 + random.random()))
    return None


def _read_capped(zf, info, max_bytes):
    """
    Read a zip member without trusting its declared size: decompress at most `max_bytes`
    and reject (return None) if the stream is larger. Defends against compression bombs
    and spoofed central-directory sizes.
    """
    try:
        with zf.open(info) as fh:
            data = fh.read(max_bytes + 1)
        return None if len(data) > max_bytes else data
    except Exception:
        return None


# ---------------------------------------------------------------------------
# ID resolution + retrieval
# ---------------------------------------------------------------------------
def pmid_to_pmcid(pmid) -> Optional[str]:
    """Resolve a PMID to a PMCID (e.g. 'PMC1234567') via the NCBI ID Converter, or None."""
    try:
        resp = _http_get(ID_CONVERTER_URL, params=_ncbi_params({"ids": str(pmid), "format": "json"}))
        for rec in (resp.json() or {}).get("records", []):
            if rec.get("pmcid"):
                return rec["pmcid"]
    except Exception:
        return None
    return None


def _normalize_pmcid(pmid) -> Optional[str]:
    s = str(pmid).strip()
    if s.upper().startswith("PMC"):
        return s.upper()
    return pmid_to_pmcid(s)


def fetch_supplementary_zip(pmcid, *, caps) -> Optional[zipfile.ZipFile]:
    """Fetch the Europe PMC supplementaryFiles ZIP into memory (size-capped stream, cached by
    pmcid). None if absent / too big / not a zip. A fresh ZipFile is returned each call (over the
    cached bytes) since a ZipFile stream is single-use."""
    with _cache_lock:
        hit = pmcid in _ZIP_CACHE
        if hit:
            _ZIP_CACHE.move_to_end(pmcid)
            content = _ZIP_CACHE[pmcid]
    if not hit:
        content = _download_capped(f"{EPMC_BASE}/{pmcid}/supplementaryFiles", caps["max_zip_bytes"])
        _cache_put(_ZIP_CACHE, pmcid, content)
    if not content or content[:2] != b"PK":
        return None
    try:
        return zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile:
        return None


def fetch_supplementary_captions(pmcid) -> Dict[str, str]:
    """Map supplementary filename -> 'label: caption' parsed from the JATS fullTextXML (cached)."""
    with _cache_lock:
        if pmcid in _CAPTION_CACHE:
            _CAPTION_CACHE.move_to_end(pmcid)
            return _CAPTION_CACHE[pmcid]
    out: Dict[str, str] = {}
    try:
        resp = _http_get(f"{EPMC_BASE}/{pmcid}/fullTextXML")
        from lxml import etree
        root = etree.fromstring(resp.content)
        for sm in root.iter():
            if not isinstance(sm.tag, str) or etree.QName(sm).localname != "supplementary-material":
                continue
            href = sm.get(XLINK)
            label = caption = ""
            for child in sm.iter():
                ln = etree.QName(child).localname
                if ln == "media" and not href:
                    href = child.get(XLINK)
                elif ln == "label" and not label:
                    label = " ".join("".join(child.itertext()).split())
                elif ln == "caption" and not caption:
                    caption = " ".join("".join(child.itertext()).split())
            if not href:
                continue
            key = os.path.basename(href)
            out[key] = " — ".join(x for x in (label, caption) if x)[:500]
    except Exception:
        return _cache_put(_CAPTION_CACHE, pmcid, out)
    return _cache_put(_CAPTION_CACHE, pmcid, out)


def _match_caption(name, captions) -> str:
    if name in captions:
        return captions[name]
    stem = os.path.splitext(name)[0].lower()
    for k, v in captions.items():
        if os.path.splitext(k)[0].lower() == stem:
            return v
    return ""


# ---------------------------------------------------------------------------
# Gene-aware matching
# ---------------------------------------------------------------------------
def _build_gene_regex(gene_id, aliases):
    """Boundary-aware regex over the gene id + aliases (tolerant of hyphen/underscore/space)."""
    terms = [str(t) for t in ([gene_id] + list(aliases or [])) if t and len(str(t)) >= 3]
    pats = []
    for t in terms:
        m = re.fullmatch(r"([A-Za-z]+)([0-9]+)", t)
        if m:
            letters, digits = m.groups()
            inner = rf"{re.escape(letters)}-?{re.escape(digits)}"
        else:
            parts, i = [], 0
            while i < len(t):
                if t[i] in "_- ":
                    while i < len(t) and t[i] in "_- ":
                        i += 1
                    parts.append(r"[-_\s]+")
                    continue
                parts.append(re.escape(t[i]))
                i += 1
            inner = "".join(parts)
        pats.append(inner)
    if not pats:
        return None
    return re.compile(rf"(?<![A-Za-z0-9])(?:{'|'.join(pats)})(?![A-Za-z0-9])", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Decoding + per-format parsers
# ---------------------------------------------------------------------------
def _decode(data: bytes) -> str:
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(data).best()
        if best:
            return str(best)
    except Exception:
        pass
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_pdf(data, caps):
    import pdfplumber
    candidates = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for pno, page in enumerate(pdf.pages, 1):
            for line in (page.extract_text() or "").splitlines():
                line = line.strip()
                if line:
                    candidates.append(f"[p{pno}] {line}")
            for tbl in (page.extract_tables() or []):
                for row in tbl:
                    cells = [(c or "").strip() for c in row]
                    if any(cells):
                        candidates.append(f"[p{pno} table] " + " | ".join(cells))
    return {"kind": "text", "always": [], "candidates": candidates}


def _parse_spreadsheet(data, ext, caps):
    import pandas as pd
    engine = "openpyxl" if ext == ".xlsx" else "xlrd"
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, engine=engine, dtype=str, header=0)
    always, candidates = [], []
    for sheet, df in sheets.items():
        df = df.fillna("")
        always.append(f"[sheet {sheet}] columns: " + " | ".join(str(c) for c in df.columns))
        for _, row in df.iterrows():
            candidates.append(f"[{sheet}] " + " | ".join(str(x) for x in row.values))
    return {"kind": "table", "always": always, "candidates": candidates}


def _parse_delimited(data, ext, caps):
    import pandas as pd
    text = _decode(data)
    sample = text[:8192]
    delim = None
    try:
        delim = csv.Sniffer().sniff(sample, delimiters=",\t;|").delimiter
    except Exception:
        if ext in (".tsv", ".tab"):
            delim = "\t"
        elif ext == ".csv":
            delim = ","
        else:
            first = sample.splitlines()[0] if sample.splitlines() else ""
            delim = max([",", "\t", ";", "|"], key=lambda d: first.count(d)) if first else ","
    always, candidates = [], []
    try:
        df = pd.read_csv(io.StringIO(text), sep=delim, dtype=str,
                         engine="python", on_bad_lines="skip").fillna("")
        always.append("columns: " + " | ".join(str(c) for c in df.columns))
        for _, row in df.iterrows():
            candidates.append(" | ".join(str(x) for x in row.values))
    except Exception:
        lines = [l for l in text.splitlines() if l.strip()]
        if lines:
            always.append("header: " + lines[0])
            candidates.extend(lines[1:])
    return {"kind": "table", "always": always, "candidates": candidates}


def _parse_docx(data, caps):
    import docx
    doc = docx.Document(io.BytesIO(data))
    candidates = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for ti, tbl in enumerate(doc.tables, 1):
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                candidates.append(f"[table{ti}] " + " | ".join(cells))
    return {"kind": "text", "always": [], "candidates": candidates}


def _parse_markup(data, ext, caps):
    from bs4 import BeautifulSoup
    parser = "lxml-xml" if ext == ".xml" else "lxml"
    soup = BeautifulSoup(_decode(data), parser)
    candidates = [l.strip() for l in soup.get_text("\n").splitlines() if l.strip()]
    return {"kind": "text", "always": [], "candidates": candidates}


def _parse_doc(data, caps):
    """Legacy binary .doc (OLE): best-effort text extraction from the WordDocument stream.
    Pulls both UTF-16LE and 8-bit printable runs so gene ids/rows survive (formatting is lost)."""
    import olefile
    bio = io.BytesIO(data)
    if not olefile.isOleFile(bio):
        return {"kind": "skipped", "always": [], "candidates": [], "note": "not an OLE .doc"}
    ole = olefile.OleFileIO(bio)
    try:
        if not ole.exists("WordDocument"):
            return {"kind": "skipped", "always": [], "candidates": [], "note": ".doc without WordDocument stream"}
        raw = ole.openstream("WordDocument").read()
    finally:
        ole.close()

    candidates, seen = [], set()

    def _add(s):
        s = s.strip()
        if len(s) >= 4 and s not in seen:
            seen.add(s)
            candidates.append(s)

    for m in re.findall(rb"(?:[\x20-\x7e]\x00){4,}", raw):   # UTF-16LE runs (Word's usual encoding)
        _add(m.decode("utf-16le", errors="ignore"))
    for m in re.findall(rb"[\x20-\x7e]{4,}", raw):           # 8-bit / cp1252 runs
        _add(m.decode("cp1252", errors="ignore"))
    return {"kind": "text", "always": [], "candidates": candidates}


def _parse_entry(name, data, caps):
    ext = os.path.splitext(name)[1].lower()
    try:
        if ext == ".pdf":
            return _parse_pdf(data, caps)
        if ext == ".xlsx":
            return _parse_spreadsheet(data, ".xlsx", caps)
        if ext == ".xls":
            return _parse_spreadsheet(data, ".xls", caps)
        if ext in (".csv", ".tsv", ".tab", ".txt"):
            return _parse_delimited(data, ext, caps)
        if ext == ".doc":
            return _parse_doc(data, caps)
        if ext == ".docx":
            return _parse_docx(data, caps)
        if ext in (".xml", ".html", ".htm"):
            return _parse_markup(data, ext, caps)
        if ext in IMAGE_EXTS:
            return {"kind": "image", "always": [], "candidates": [], "note": "image (deferred)"}
    except Exception as e:
        return {"kind": "error", "always": [], "candidates": [], "note": f"parse error: {type(e).__name__}"}
    return {"kind": "skipped", "always": [], "candidates": [], "note": f"unsupported ext {ext or '(none)'}"}


# ---------------------------------------------------------------------------
# Assembly
# ---------------------------------------------------------------------------
def _assemble_file_block(name, parsed, captions, gene_re, caps) -> str:
    caption = _match_caption(name, captions)
    header_line = f"[{name}]" + (f" — {caption}" if caption else "")

    kind = parsed.get("kind")
    always = list(parsed.get("always", []))[: caps["max_always"]]
    candidates = parsed.get("candidates", []) or []

    matched = []
    if gene_re is not None:
        for rec in candidates:
            if gene_re.search(rec):
                matched.append(rec)
                if len(matched) >= caps["max_matched_rows"]:
                    break

    body = list(always)
    if matched:
        body.extend(matched)
    elif kind == "table":
        # the caption / column line already describes the table; skip a row preview
        if not body:
            return ""
    else:
        preview = candidates[: caps["preview_rows"]]
        if not body and not preview:
            return ""
        if preview:
            body.append("(gene not found in this file; first lines:)")
            body.extend(preview)

    text = header_line + "\n" + "\n".join(body)
    if len(text) > caps["max_file_chars"]:
        text = text[: caps["max_file_chars"]] + " …[truncated]"
    return text


def _resolve_aliases(gene_id, aliases, host_db):
    if aliases is None and host_db:
        try:
            from pipeline.vpdb_helpers import get_gene_synonyms
            aliases = get_gene_synonyms(gene_id, "", host_db)
        except Exception:
            aliases = []
    return list(aliases or [])


def _new_state():
    return {"files_seen": 0, "extracted_bytes": 0, "image_count": 0,
            "other_files": [], "truncated": False}


def _walk_zip(zf, caps, state, depth, handle_file):
    """
    Traverse a zip (recursing up to `max_zip_depth` into nested zips) applying all size guards,
    and hand each parsed, non-image, non-empty file to `handle_file(name, parsed, state)`.
    `handle_file` returns True to stop the whole traversal (e.g. output budget reached).
    Returns True if traversal was stopped early.
    """
    for info in zf.infolist():
        if state["files_seen"] >= caps["max_files"] or state["extracted_bytes"] >= caps["max_total_extract_bytes"]:
            state["truncated"] = True
            return True
        if info.is_dir():
            continue
        name = os.path.basename(info.filename)
        if not name:
            continue
        state["files_seen"] += 1
        if info.file_size > caps["max_file_bytes"]:
            state["other_files"].append(f"{name} (too large, skipped)")
            continue
        remaining = caps["max_total_extract_bytes"] - state["extracted_bytes"]
        data = _read_capped(zf, info, min(caps["max_file_bytes"], remaining))
        if data is None:
            state["other_files"].append(f"{name} (over size cap)")
            continue
        state["extracted_bytes"] += len(data)

        ext = os.path.splitext(name)[1].lower()
        if ext == ".zip":
            if depth >= caps["max_zip_depth"]:
                state["other_files"].append(f"{name} (nested zip, depth limit)")
                continue
            try:
                nested = zipfile.ZipFile(io.BytesIO(data))
            except zipfile.BadZipFile:
                state["other_files"].append(name)
                continue
            if _walk_zip(nested, caps, state, depth + 1, handle_file):
                return True
            continue

        parsed = _parse_entry(name, data, caps)
        kind = parsed.get("kind")
        if kind == "image":
            state["image_count"] += 1
            continue
        if kind in ("skipped", "error") and not parsed.get("candidates") and not parsed.get("always"):
            state["other_files"].append(name)  # unparseable data file (e.g. .pptx)
            continue
        if handle_file(name, parsed, state):
            return True
    return False


def _parsed_files(pmcid, caps):
    """Parse a paper's supplement ONCE and cache the result by pmcid, so a paper with N genes
    parses its (possibly large) files once instead of N times. Returns
    {"files":[(name,parsed)], "image_count", "other_files"} or None if there is no supplement."""
    with _cache_lock:
        if pmcid in _PARSED_CACHE:
            _PARSED_CACHE.move_to_end(pmcid)
            return _PARSED_CACHE[pmcid]
    zf = fetch_supplementary_zip(pmcid, caps=caps)
    if zf is None:
        _cache_put(_PARSED_CACHE, pmcid, None)
        return None
    files = []
    state = _new_state()

    def _grab(name, parsed, st):
        files.append((name, parsed))
        return False

    try:
        _walk_zip(zf, caps, state, 0, _grab)
    except Exception:
        pass
    result = {"files": files, "image_count": state["image_count"], "other_files": state["other_files"]}
    _cache_put(_PARSED_CACHE, pmcid, result)
    return result


def get_supplementary_text(pmid, gene_id, aliases=None, host_db=None, *, caps=None) -> str:
    """
    Return a bounded, gene-filtered supplementary-materials text block for (pmid, gene_id),
    or "" if there are no supplements / the paper is not resolvable / anything fails.
    """
    caps = {**DEFAULT_CAPS, **(caps or {})}
    try:
        aliases = _resolve_aliases(gene_id, aliases, host_db)
        pmcid = _normalize_pmcid(pmid)
        if not pmcid:
            return ""
        pf = _parsed_files(pmcid, caps)   # parsed once per paper, cached
        if pf is None:
            return ""

        captions = fetch_supplementary_captions(pmcid)  # best-effort enrichment
        gene_re = _build_gene_regex(gene_id, aliases)

        blocks, total, truncated = [], 0, False
        for name, parsed in pf["files"]:
            block = _assemble_file_block(name, parsed, captions, gene_re, caps)
            if not block:
                continue
            if total + len(block) > caps["total_char_budget"]:
                truncated = True
                break
            blocks.append(block)
            total += len(block)

        if not blocks:
            return ""  # only images / unparseable files — nothing text-extractable

        extras = []
        if pf["image_count"]:
            extras.append(f"[{pf['image_count']} figure/image file(s) present — image extraction deferred]")
        if pf["other_files"]:
            shown = ", ".join(pf["other_files"][:8])
            more = f" (+{len(pf['other_files']) - 8} more)" if len(pf["other_files"]) > 8 else ""
            extras.append(f"[{len(pf['other_files'])} non-text file(s) not parsed: {shown}{more}]")
        if extras:
            blocks.append("\n".join(extras))

        head = f"=== SUPPLEMENTARY MATERIALS (gene {gene_id}"
        if aliases:
            head += " / " + ", ".join(aliases[:5])
        head += f"; PMCID {pmcid}) ==="
        tail = "\n[supplementary truncated: character budget reached]" if truncated else ""
        return head + "\n" + "\n\n".join(blocks) + tail + "\n=== END SUPPLEMENTARY MATERIALS ===\n"
    except Exception:
        return ""


def count_supplementary_mentions(pmid, gene_id, aliases=None, host_db=None, *, caps=None) -> dict:
    """
    Software-only scan (no LLM): does the gene id / any alias appear anywhere in the paper's
    supplementary files (incl. nested zips)? Used to measure how many main-text-negative pairs
    get 'unlocked' by supplements. Returns:
        {'pmcid', 'available' (a suppl zip existed), 'found', 'mentions', 'files': [(name, n), ...]}
    Never raises.
    """
    caps = {**DEFAULT_CAPS, **(caps or {})}
    out = {"pmcid": None, "available": False, "found": False, "mentions": 0, "files": []}
    try:
        aliases = _resolve_aliases(gene_id, aliases, host_db)
        pmcid = _normalize_pmcid(pmid)
        out["pmcid"] = pmcid
        if not pmcid:
            return out
        pf = _parsed_files(pmcid, caps)   # parsed once per paper, cached
        if pf is None:
            return out
        out["available"] = True
        gene_re = _build_gene_regex(gene_id, aliases)
        if gene_re is None:
            return out

        mentions, files = 0, []
        for name, parsed in pf["files"]:
            blob = "\n".join(list(parsed.get("always", [])) + list(parsed.get("candidates", [])))
            n = len(gene_re.findall(blob)) if blob else 0
            if n:
                mentions += n
                files.append((name, n))
        out["mentions"] = mentions
        out["files"] = files
        out["found"] = mentions > 0
        return out
    except Exception:
        return out
